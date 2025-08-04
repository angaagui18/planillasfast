import streamlit as st
import pandas as pd
import pdfplumber
import zipfile
import os
import tempfile
from docx import Document
from fuzzywuzzy import fuzz
import datetime
import unidecode

# ------------------ FUNCIONES ------------------ #

def leer_contrato(archivo_pdf):
    texto_total = ""
    with pdfplumber.open(archivo_pdf) as pdf:
        for pagina in pdf.pages:
            texto_total += pagina.extract_text() + "\n"
    return texto_total

def leer_planillas(zip_file):
    planillas = []
    with zipfile.ZipFile(zip_file, 'r') as z:
        for nombre_archivo in z.namelist():
            if nombre_archivo.endswith('.xlsx'):
                with z.open(nombre_archivo) as file:
                    df = pd.read_excel(file)
                    df["nombre_archivo"] = nombre_archivo
                    planillas.append(df)
    return planillas

def generar_informe_word(texto_contrato, planillas, nombre_output):
    doc = Document()
    doc.add_heading("Informe de revisión de planillas", level=1)

    for i, df in enumerate(planillas):
        doc.add_heading(f"Planilla: {df['nombre_archivo'][0]}", level=2)
        doc.add_paragraph(f"Filas: {len(df)} columnas: {len(df.columns)}")
        
        # Aquí va tu lógica personalizada de verificación
        columnas_contrato = [col.lower() for col in texto_contrato.split()]
        columnas_planilla = [col.lower() for col in df.columns]

        for columna in columnas_planilla:
            coincidencia = max([fuzz.ratio(columna, ref) for ref in columnas_contrato])
            if coincidencia < 50:
                doc.add_paragraph(f"⚠️ Posible inconsistencia: '{columna}' no se encuentra en el contrato.", style='List Bullet')

    doc.save(nombre_output)

# ------------------ INTERFAZ ------------------ #

st.set_page_config(page_title="PlanillasFast", layout="centered")
st.title("📑 PlanillasFast - Revisión de contratos y planillas")

st.info("Sube un contrato (PDF) y un ZIP con las planillas (Excel). El sistema generará un informe Word con los hallazgos.")

contrato = st.file_uploader("📄 Sube el contrato (PDF)", type=["pdf"])
zip_planillas = st.file_uploader("📁 Sube las planillas (ZIP con Excel)", type=["zip"])

if st.button("Generar informe"):
    if contrato is not None and zip_planillas is not None:
        with st.spinner("Procesando archivos..."):
            texto = leer_contrato(contrato)
            planillas = leer_planillas(zip_planillas)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmpfile:
                generar_informe_word(texto, planillas, tmpfile.name)
                st.success("✅ Informe generado con éxito.")
                with open(tmpfile.name, "rb") as f:
                    st.download_button("📥 Descargar informe Word", f, file_name="informe_planillas.docx")
    else:
        st.warning("⚠️ Por favor, sube ambos archivos para continuar.")
