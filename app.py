import streamlit as st
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO

st.set_page_config(page_title="PlanillasFast", layout="centered")

st.title("📄 PlanillasFast")
st.info(
    "Sube un contrato (PDF) y una planilla (Excel). El sistema generará un informe Word con los hallazgos."
)

# Subida del contrato
contrato_pdf = st.file_uploader("📑 Sube el contrato (PDF)", type=["pdf"])

# Subida de la planilla Excel
planilla_excel = st.file_uploader("📊 Sube la planilla (.xlsx)", type=["xlsx"])

# Función para analizar PDF (contrato)
def extraer_texto_contrato(file):
    try:
        reader = PdfReader(file)
        texto = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                texto += page_text
        return texto.strip()
    except Exception as e:
        st.error(f"❌ Error al leer el contrato PDF: {e}")
        return ""

# Función para generar informe Word
def generar_informe(texto_contrato, planilla_df):
    try:
        doc = Document()
        doc.add_heading("Informe de Revisión de Planillas", 0)

        doc.add_heading("Contrato", level=1)
        doc.add_paragraph(
            texto_contrato[:1500] + "..." if texto_contrato else "No se pudo extraer texto del contrato."
        )

        doc.add_heading("Resumen de la Planilla", level=1)
        doc.add_paragraph(f"Filas totales en la planilla: {len(planilla_df)}")

        columnas_planilla = [str(col).lower() for col in planilla_df.columns]
        doc.add_paragraph("Columnas encontradas en la planilla:")
        doc.add_paragraph(", ".join(columnas_planilla))

        doc.add_heading("Primeras 5 filas", level=2)
        doc.add_paragraph(planilla_df.head().to_string())

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"❌ Error al generar el informe: {e}")
        return None

# Botón de generación
if st.button("Generar informe"):
    if contrato_pdf and planilla_excel:
        try:
            texto_contrato = extraer_texto_contrato(contrato_pdf)
            planilla_df = pd.read_excel(planilla_excel)

            informe = generar_informe(texto_contrato, planilla_df)

            if informe:
                st.success("✅ Informe generado correctamente.")
                st.download_button(
                    label="📥 Descargar informe Word",
                    data=informe,
                    file_name="informe_planillas.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"❌ Error al procesar los archivos: {e}")
    else:
        st.warning("⚠️ Por favor, sube tanto el contrato como la planilla.")
